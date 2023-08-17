import requests
from bs4 import BeautifulSoup
from docx import Document

# Define the URL of the website
website_url = 'https://learn.codespace.co.za/'

# Define a function to scrape the content of a URL
def scrape_url(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')
        return soup.get_text()
    else:
        return None

# Replace these URLs with the actual URLs from the dropdown
dropdown_urls = [
    "https://learn.codespace.co.za/courses/144",
    "https://learn.codespace.co.za/courses/142",
    "https://learn.codespace.co.za/courses/147",
    # Add more URLs as needed
]

# Create a new Word document
doc = Document()

# Scrape and add the content of each URL to the Word document
for url in dropdown_urls:
    content = scrape_url(url)
    if content is not None:
        doc.add_heading(f"Content of {url}", level=1)
        doc.add_paragraph(content)
        doc.add_page_break()
    else:
        print(f"Failed to retrieve content from {url}")

# Save the Word document
doc.save("scraped_content.docx")
